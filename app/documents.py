from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


SUPPORTED_SUFFIXES = {".pdf", ".docx", ".md", ".markdown", ".txt"}


@dataclass(frozen=True)
class SourceDocument:
    source_name: str
    location: str
    text: str
    metadata: dict = field(default_factory=dict)


def parse_document(filename: str, content: bytes, metadata: dict | None = None) -> list[SourceDocument]:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"不支持的文件类型：{suffix or '无扩展名'}")

    if suffix == ".pdf":
        reader = PdfReader(BytesIO(content))
        documents = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                documents.append(SourceDocument(filename, f"第 {page_number} 页", text, metadata or {}))
        return documents

    if suffix == ".docx":
        document = Document(BytesIO(content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    else:
        text = content.decode("utf-8", errors="replace")

    return [SourceDocument(filename, "全文", text, metadata or {})] if text.strip() else []
